#!/usr/bin/env python3
"""
构建邮政金融业务知识文档向量数据库

将 postal-knowledge-qa 技能目录下的 .docx 知识文档（如 知识库.docx、
金融百度相关知识库.docx）抽取文本、按章节分块，用本地 bge-m3 嵌入后
存入 ChromaDB 的独立集合 postal_knowledge_docs（持久化到 deer-flow/data/chroma）。

用法：
    python scripts/build_doc_vector_db.py [--rebuild]
"""
import argparse
import re
from pathlib import Path

import docx
import requests

# 路径配置
PROJECT_ROOT = Path(__file__).resolve().parents[2]  # PostViewAgent/
SKILL_DIR = PROJECT_ROOT / "deer-flow" / "skills" / "public" / "postal-knowledge-qa"

# ChromaDB 持久化目录（与规则库一致）
CHROMA_DIR = PROJECT_ROOT / "deer-flow" / "data" / "chroma"
COLLECTION_NAME = "postal_knowledge_docs"

MAX_CHARS = 400  # 每个知识块的目标字符数

OLLAMA_URL = "http://192.168.7.88:11434"
MODEL = "bge-m3:latest"


class BgeM3EmbeddingFunction:
    """本地 Ollama bge-m3 嵌入函数，兼容 ChromaDB EmbeddingFunction 接口"""

    def __init__(self, url: str = OLLAMA_URL, model: str = MODEL):
        self.url = f"{url}/api/embed"
        self.model = model

    def _embed(self, texts):
        if isinstance(texts, str):
            texts = [texts]
        if texts and isinstance(texts[0], list):
            texts = texts[0]
        resp = requests.post(
            self.url,
            json={"model": self.model, "input": list(texts)},
            timeout=300,
        )
        resp.raise_for_status()
        return resp.json()["embeddings"]

    def __call__(self, input):
        return self._embed(input)

    def embed_query(self, input):
        return self._embed(input)[0]

    def embed_documents(self, input):
        return self._embed(input)

    def name(self):
        return "bge-m3"


def extract_text(docx_path: Path) -> list[str]:
    """抽取 docx 全部段落与表格文本，返回非空行列表"""
    d = docx.Document(str(docx_path))
    lines = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                lines.append(line)
    return lines


_HEADING_RE = re.compile(
    r"^(第[一二三四五六七八九十百]+[章节部分]|[一二三四五六七八九十]+、|"
    r"\d+[\.、．]|\d+\.\d+[\.、．]?|（[一二三四五六七八九十]+）|\([0-9]+\)|"
    r"[A-Z][\.、．])"
)


def _is_heading(line: str) -> bool:
    """判断一行是否像章节标题：较短且以编号/序号开头"""
    if len(line) > 25:
        return False
    if line.endswith(("。", "；", "，", "：")):
        return False
    return bool(_HEADING_RE.match(line))


def chunk_lines(lines: list[str], max_chars: int = MAX_CHARS) -> list[dict]:
    """按章节分块：累积段落至目标长度，并用最近的小节标题作为上下文前缀"""
    blocks = []
    current = []
    current_len = 0
    section_title = ""
    block_no = 0

    def flush():
        nonlocal current, current_len, block_no
        if not current:
            return
        text = "\n".join(current)
        if section_title and not text.startswith(section_title):
            text = f"{section_title}\n{text}"
        blocks.append({
            "id": f"doc_{block_no}",
            "text": text,
            "title": section_title or (current[0] if current else ""),
        })
        block_no += 1
        current = []
        current_len = 0

    for line in lines:
        if _is_heading(line):
            # 新小节开始：先冲刷上一块，再更新小节标题
            flush()
            section_title = line
            continue
        current.append(line)
        current_len += len(line)
        if current_len >= max_chars:
            flush()

    flush()
    return blocks


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--rebuild", action="store_true", help="重建集合（清空旧数据）")
    args = parser.parse_args()

    import chromadb

    CHROMA_DIR.mkdir(parents=True, exist_ok=True)
    client = chromadb.PersistentClient(path=str(CHROMA_DIR))
    ef = BgeM3EmbeddingFunction()

    collection = client.get_or_create_collection(
        COLLECTION_NAME,
        metadata={"hnsw:space": "cosine"},
        embedding_function=ef,
    )

    if args.rebuild:
        print("[INFO] 重建模式：清空旧集合")
        try:
            client.delete_collection(COLLECTION_NAME)
        except Exception:
            pass
        collection = client.create_collection(
            COLLECTION_NAME,
            metadata={"hnsw:space": "cosine"},
            embedding_function=ef,
        )

    # 收集所有文档块（ID 加文档前缀避免跨文件冲突）
    blocks = []
    for f in sorted(SKILL_DIR.glob("*.docx")):
        lines = extract_text(f)
        chunks = chunk_lines(lines)
        print(f"[INFO] {f.name}: {len(lines)} 行 -> {len(chunks)} 个知识块")
        prefix = f.stem
        for c in chunks:
            c["id"] = f"{prefix}_{c['id']}"
            c["metadata"] = {
                "type": "document",
                "source": f.name,
                "title": c.pop("title"),
            }
            blocks.append(c)

    print(f"[INFO] 共加载 {len(blocks)} 个知识块")

    # 分批写入（避免一次请求过大，bge-m3 长文本嵌入较慢）
    import time

    batch_size = 5
    total = len(blocks)
    t_start = time.time()
    for i in range(0, total, batch_size):
        batch = blocks[i : i + batch_size]
        ids = [b["id"] for b in batch]
        docs = [b["text"] for b in batch]
        metas = [b["metadata"] for b in batch]
        t0 = time.time()
        collection.upsert(ids=ids, documents=docs, metadatas=metas)
        elapsed = time.time() - t0
        done = min(i + len(batch), total)
        pct = done / total * 100
        print(
            f"  [{done}/{total} {pct:.0f}%] 本批 {len(batch)} 条耗时 {elapsed:.1f}s "
            f"(累计 {time.time()-t_start:.1f}s)"
        )

    count = collection.count()
    print(f"\n[OK] 文档向量库构建完成，共 {count} 条记录")
    print(f"    存储位置: {CHROMA_DIR}")


if __name__ == "__main__":
    main()
